from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT

# 1) Criação do documento
doc = Document()

# 2) Configuração para LANDSCAPE (Paisagem)
section = doc.sections[0]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

# Margens
section.left_margin = Cm(1.27)
section.right_margin = Cm(1.27)
section.top_margin = Cm(1.27)
section.bottom_margin = Cm(1.27)

# Título (ajustado para 2026/2027)
doc.add_heading('Calendário Integrado: Governança & Safra (2026/2027)', 0)

# 3) Pauta fixa mensal (incorpora anotação #4)
PAUTA_FIXA_MENSAL = (
    "• Acompanhamento da Dívida e captações.\n"
    "• Projeção de fluxo de caixa.\n"
    "• Acompanhamento mensal dos projetos de investimentos aprovados "
    "(cronograma de execução e retorno realizado x esperado).\n"
)

def com_pauta_fixa(pautas: str) -> str:
    return PAUTA_FIXA_MENSAL + pautas

# 4) Dados consolidados (Abril -> Março) com anotações incorporadas (#1, #2, #3, #4)
data = [
    ("Abril", "Estratégia Corporativa e Cenário Financeiro",
     com_pauta_fixa(
         "• Políticas de Gestão de Risco.\n"
         "• Fluxo de caixa e Endividamento.\n"
         "• Definir evento anual de liderança.\n"
         "• Evoluções na planilha de controle: apresentação de Hedges e Comercialização de Gado."
     ),
     "AVB; Pessoas; Fernanda; José Carlos; Laguna; Beatriz;",
     "• [Pau Dalho] Cana Safra: Início Colheita (1ª quinzena)\n"
     "• [Pau Dalho] Cana Plantio: Fim (1ª quinzena)\n"
     "• [Canadá] Cana Safra: Início Colheita (1ª quinzena)\n"
     "• [Canadá] Cana Plantio: Fim (1ª quinzena)\n"
     "• [Rancho Alegre] Soja: Fim Colheita (1ª Quinzena)"),

    ("Maio", "Apresentação de Resultados (Safra anterior)",
     com_pauta_fixa(
         "• Políticas de Gestão de Risco.\n"
         "• Resultados operacionais por unidade e consolidado, safra anterior.\n"
         "• Discutir e formalizar a política de comercialização de Gado (versão final)."
     ),
     "AVB; Laguna; Frank; Claudia; Beatriz; José Neto;",
     "-"),

    ("Junho", "Feedback implantação sistema Senior",
     com_pauta_fixa(
         "• Políticas de Gestão de Risco.\n"
         "• Feedback migração ERP Senior.\n"
         "• Planejamento estratégico. Cenário econômico safra 26/27 para eventuais revisões de orçamento e fluxo de caixa.\n"
         "• Diretrizes iniciais para desenvolvimento da política de comercialização de Cana (escopo, responsáveis e cronograma)."
     ),
     "AVB; Fernanda; Claudia; Frank; José Carlos; Helcio; Mayke; Beatriz;",
     "• [São Judas] Safrinha Milho: Fim Colheita (2ª Quinzena)\n"
     "• [Rancho Alegre] Safrinha Milho: Fim Colheita (2ª Quinzena)"),

    ("Julho", "Painel de Indicadores",
     com_pauta_fixa(
         "• Políticas de Gestão de Risco.\n"
         "• Painel de indicadores (ROA, ROIC, TIR)."
     ),
     "ECOWA; Fernanda; Laguna; Beatriz;",
     "-"),

    ("Agosto", "Encerramento KPMG e Controle Interno",
     com_pauta_fixa(
         "• Resultados Q1 (Abr-Jun) / e-mail pacote.\n"
         "• Políticas de Gestão de Risco.\n"
         "• Apresentação KPMG e Controle interno.\n"
         "• Feedback migração ERP Senior."
     ),
     "AVB; Laguna; José Neto; Beatriz; Frank; Claudia; KPMG;",
     "-"),

    ("Setembro", "Planejamento estratégico",
     com_pauta_fixa(
         "• Políticas de Gestão de Risco (foco exclusivo).\n"
         "Obs: Monitoramento de políticas internas.\n"
         "• Planejamento estratégico, revisão do plurianual.\n"
         "• Revisão do cronograma anual de reuniões para detalhar marcos operacionais "
         "(inclui análise do plantio da Soja em dezembro).\n"
         "• Política de comercialização de Cana: proposta inicial e plano de evolução."
     ),
     "Fernanda; Laguna; Beatriz;",
     "• [São Judas] Milho: Início Plantio (1ª Quinzena)"),

    ("Outubro", "Espaço para Demandas Especiais (ad hoc)",
     com_pauta_fixa(
         "• Demandas emergenciais.\n"
         "• Políticas de Gestão de Risco.\n"
         "Obs: Revisões excepcionais."
     ),
     "Laguna; Beatriz;",
     "• [Rancho Alegre] Soja: Início Plantio (1ª Quinzena)\n"
     "• [São Judas] Soja: Início Plantio (1ª Quinzena)\n"
     "• [Pau Dalho] Cana Safra: Fim Colheita (2ª quinzena)\n"
     "• [Pau Dalho] Soja: Início Plantio (2ª quinzena)\n"
     "• [Pau Dalho] Milho: Início Plantio (2ª quinzena)\n"
     "• [Canadá] Cana Safra: Fim Colheita (2ª quinzena)"),

    ("Novembro", "Avaliação Trimestral de Desempenho Q2",
     com_pauta_fixa(
         "• Resultados Q2 (Jul-Set) / e-mail pacote.\n"
         "• Políticas de Gestão de Risco."
     ),
     "AVB; Laguna; Beatriz; José Neto;",
     "• [Canadá] Soja: Início Plantio (1ª Quinzena)"),

    ("Dezembro", "Projeções do Ano Safra",
     com_pauta_fixa(
         "• Simulações fechamento ano safra.\n"
         "• Políticas de Gestão de Risco.\n"
         "• Projeção para CAPEX de fevereiro.\n"
         "• Planejamento estratégico. Cenário econômico safra 26/27 para eventuais revisões de orçamento e fluxo de caixa.\n"
         "• Análise do plantio da Soja (status, aderência ao plano, riscos e lições aprendidas)."
     ),
     "AVB; Fernanda; Laguna; Beatriz; José Neto;",
     "-"),

    ("Janeiro", "Recesso",
     "• Período de recesso, sem reuniões agendadas.\n"
     "Obs: Período reservado para recesso institucional.",
     "-",
     "• [São Judas] Safrinha Milho: Início Plantio (1ª Quinzena)\n"
     "• [São Judas] Soja: Fim Colheita (2ª Quinzena)\n"
     "• [São Judas] Milho: Fim Colheita (2ª Quinzena)"),

    ("Fevereiro", "Planejamento de Investimentos e Estrutura de Capital",
     com_pauta_fixa(
         "• Apresentação CAPEX próxima safra.\n"
         "• Estrutura de capital.\n"
         "• Políticas de Gestão de Risco (mensal).\n"
         "Obs: Base para definição de metas em março.\n"
         "• Resultados Q3 (Out-Dez) / e-mail pacote.\n"
         "• Planejamento estratégico. Cenário econômico safra 26/27 para eventuais revisões de orçamento e fluxo de caixa."
     ),
     "AVB; MOP; ECOWA; Laguna; Beatriz; José Neto;",
     "• [Pau Dalho] Cana Plantio: Início (1ª quinzena)\n"
     "• [Pau Dalho] Soja: Fim Colheita (1ª Quinzena)\n"
     "• [Pau Dalho] Milho: Fim Colheita (2ª Quinzena)\n"
     "• [Canadá] Cana Plantio: Início (2ª quinzena)\n"
     "• [Canadá] Soja: Fim Colheita (2ª Quinzena)\n"
     "• [Rancho Alegre] Safrinha Milho: Início Plantio (2ª Quinzena)"),

    ("Março", "Definição de Metas e Aprovação Orçamentária",
     com_pauta_fixa(
         "• Aprovação do Orçamento Operacional.\n"
         "• Estabelecimento das metas das Unidades.\n"
         "• Plano de comunicação das metas.\n"
         "• Políticas de Gestão de Risco.\n"
         "Obs: Metas monitoradas trimestralmente."
     ),
     "AVB; MOP; Laguna; Beatriz; Frank; Claudia; José Neto;",
     "-"),
]

# 5) Criação da Tabela
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# Cabeçalho
hdr_cells = table.rows[0].cells
headers = ['Mês', 'Foco (Conselho)', 'Pautas e Entregas', 'Responsável', 'Atividades Safra']
for i, text in enumerate(headers):
    paragraph = hdr_cells[i].paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = True

# Preenchimento
for mes, foco, pautas, resp, safra in data:
    row_cells = table.add_row().cells
    row_cells[0].text = mes
    row_cells[1].text = foco
    row_cells[2].text = pautas
    row_cells[3].text = resp
    row_cells[4].text = safra

# 6) Salvar o arquivo
file_name = 'Calendario_Safra_Conselho_Landscape.docx'
doc.save(file_name)
print(f"Arquivo '{file_name}' gerado com sucesso em formato Paisagem (Início: Abril).")